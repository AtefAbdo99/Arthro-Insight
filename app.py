import os
# Set environment variables to limit thread usage
os.environ['OMP_NUM_THREADS'] = '1'
os.environ['OPENBLAS_NUM_THREADS'] = '1'
os.environ['MKL_NUM_THREADS'] = '1'
os.environ['VECLIB_MAXIMUM_THREADS'] = '1'
os.environ['NUMEXPR_NUM_THREADS'] = '1'

from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for
import pandas as pd
import numpy as np
import joblib
import matplotlib
matplotlib.use('Agg')  # Use Agg backend to avoid threading issues
import matplotlib.pyplot as plt
import seaborn as sns
import io
from docx import Document
from docx.shared import Inches
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from functools import lru_cache

app = Flask(__name__)

# Lazy loading of models and data
preprocessor = None
top_features = None
models = {}
results_df = None
total_patients = 0

def load_preprocessor():
    global preprocessor
    if preprocessor is None:
        preprocessor = joblib.load('Results_for_App/preprocessor.pkl')
    return preprocessor

def load_top_features():
    global top_features
    if top_features is None:
        top_features = joblib.load('Results_for_App/top_features.pkl')
    return top_features

def load_model(model_name):
    if model_name not in models:
        models[model_name] = joblib.load(f'Results_for_App/{model_name.replace(" ", "_")}.pkl')
    return models[model_name]

@lru_cache(maxsize=1)
def load_results_df():
    return pd.read_csv('Results_for_App/model_results.csv', index_col=0)

def get_total_patients():
    global total_patients
    if total_patients == 0:
        total_patients = len(load_results_df())
    return total_patients

feature_categories = {
    'Above weight cut-off for age/gender group': ['No', 'Yes'],
    'Employment': ['Works for pay', 'Unpaid work for family business', 'Not working in part due to health', 'Not working other reasons'],
    'Education': ['Less than high school graduate', 'High school graduate', 'Some college', 'College graduate', 'Some graduate school', 'Graduate degree'],
    'Marital Status': ['Married', 'Widowed', 'Divorced', 'Separated', 'Never married'],
    'Limits or avoids activities due to pain, aching or stiffness': ['No Limits or avoidance', 'Avoids', 'Limits'],
    'Either knee, ever injured so badly difficult to walk for at least one week': ['No', 'Yes'],
    'Either knee, used medication for pain, aching or stiffness, past 12 months': ['No', 'Yes'],
    'Ever had heart attack': ['No', 'Yes'],
    'Ever treated for heart failure': ['No', 'Yes'],
    'Had operation to unclog or bypass arteries in legs': ['No', 'Yes'],
    'Had stroke, cerebrovascular accident, blood clot or bleeding in brain, or transient ischemic attack (TIA)': ['No', 'Yes'],
    'Have asthma': ['No', 'Yes'],
    'Have emphysema, chronic bronchitis, or COPD': ['No', 'Yes'],
    'Have stomach ulcers or peptic ulcer disease': ['No', 'Yes'],
    'Have diabetes': ['No', 'Yes'],
    'Ever had problem with kidneys, poor kidney function': ['No', 'Yes'],
    'Have rheumatoid arthritis': ['No', 'Yes'],
    'Used SAMe (S-adenosylmethionine) for joint pain or arthritis more than half the days of the month, past 30 days': ['No', 'Yes'],
    'Used MSM (methylsulfonylmethane) for joint pain or arthritis more than half the days of the month, past 30 days': ['No', 'Yes'],
    'Used glucosamine for joint pain or arthritis, past 6 months': ['No', 'Yes'],
    'Used parathyroid hormone or PTH (includes Forteo, or teriparatide, given by injection) to treat osteoporosis, past 6 months': ['No', 'Yes'],
    'Taken bisphosphonate medication (includes alendronate, risedronate...) to treat osteoporosis or Paget\'s disease, past 5 years': ['No', 'Yes'],
    'Analgesic use': ['Not used in last 30 days', 'Used in last 30 days'],
    'Repeated chair stands: trial 1': ['Completes 5 stands without using arms', 'Stands using arms', 'Attempted, unable to complete', 'Not attempted, unable'],
    'Repeated chair stands: trial 2': ['Completes 5 stands without using arms', 'Stands using arms', 'Attempted, unable to complete', 'Not attempted, unable'],
    '20-meter walk: trial 1': ['Completed', 'Not attempted, unable', 'Attempted, unable to complete'],
    '20-meter walk: trial 2': ['Completed', 'Not attempted, unable', 'Attempted, unable to complete']
}

numeric_ranges = {
    'Age': range(18, 101),
    'SBP': range(90, 201),
    'DBP': range(60, 121),
    'Height': range(140, 201),
    'WEIGHT': range(40, 151),
    'BMI': range(15, 41),
    'CES-D Score': range(0, 61),
    'SF-12: physical summary scale': range(0, 101),
    'SF-12: mental summary scale': range(0, 101),
    'V00PASE': range(0, 5001),
    'Left Knee WOMAC Total Score': range(0, 101),
    'Right Knee WOMAC Total Score': range(0, 101)
}

@app.route('/')
def welcome():
    return render_template('welcome.html')

@app.route('/predict', methods=['GET', 'POST'])
def predict():
    if request.method == 'POST':
        input_data = request.form.to_dict()
        
        for key in input_data:
            if key in feature_categories:
                for cat, value in enumerate(feature_categories[key]):
                    if input_data[key] == value:
                        input_data[key] = cat
                        break
        
        input_df = pd.DataFrame([input_data])
        
        preprocessor = load_preprocessor()
        top_features = load_top_features()
        
        missing_cols = set(preprocessor.feature_names_in_) - set(input_df.columns)
        for col in missing_cols:
            input_df[col] = 0  # Default value for missing columns
        
        processed_data = preprocessor.transform(input_df)
        
        top_feature_indices = [preprocessor.feature_names_in_.tolist().index(f) for f in top_features]
        processed_data = processed_data[:, top_feature_indices]
        
        predictions = {}
        for model_name in ['Logistic Regression', 'Kernel SVM', 'Decision Tree', 'Random Forest', 'Naïve Bayes', 'XGBoost', 'NGBoost', 'LightGBM', 'CatBoost']:
            model = load_model(model_name)
            prediction = model.predict_proba(processed_data)[:, 1][0]
            predictions[model_name] = float(prediction)
        
        best_model_name = get_best_model(predictions)
        
        return render_template('results.html', results=predictions, input_data=input_df.to_dict(orient='records')[0], 
                               results_df=load_results_df(), total_patients=get_total_patients(), best_model_name=best_model_name)
    
    return render_template('predict.html', top_features=list(feature_categories.keys()) + list(numeric_ranges.keys()), feature_categories=feature_categories, numeric_ranges=numeric_ranges)

def get_best_model(predictions):
    return max(predictions, key=predictions.get)

@app.route('/contact', methods=['POST'])
def contact():
    name = request.form['name']
    email = request.form['email']
    message = request.form['message']
    
    print(f"Received contact form submission from {name} ({email}): {message}")
    
    return redirect(url_for('welcome'))

@app.route('/export_results', methods=['POST'])
def export_results():
    format = request.form.get('format')
    results = eval(request.form.get('results'))
    input_data = eval(request.form.get('input_data'))
    
    if format == 'docx':
        return export_docx(results, input_data)
    elif format == 'pdf':
        return export_pdf(results, input_data)
    elif format == 'png':
        return export_png(results, input_data)
    else:
        return jsonify({'error': 'Invalid format'}), 400

def export_docx(results, input_data):
    doc = Document()
    doc.add_heading('Arthro Insight - Prediction Results', 0)
    
    doc.add_heading('Input Data', level=1)
    for key, value in input_data.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_heading('Model Predictions', level=1)
    for model, prediction in results.items():
        doc.add_paragraph(f"{model}: {prediction:.2%}")

    img_stream = io.BytesIO()
    create_results_chart(results)
    plt.savefig(img_stream, format='png')
    img_stream.seek(0)
    doc.add_picture(img_stream, width=Inches(6))

    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    return send_file(doc_stream, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True, download_name='arthro_insight_results.docx')

def export_pdf(results, input_data):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 50, 'Arthro Insight - Prediction Results')

    c.setFont("Helvetica", 12)
    c.drawString(50, height - 80, 'Input Data:')
    y = height - 100
    for key, value in input_data.items():
        c.drawString(50, y, f"{key}: {value}")
        y -= 20

    c.drawString(50, y, 'Model Predictions:')
    y -= 20
    for model, prediction in results.items():
        c.drawString(50, y, f"{model}: {prediction:.2%}")
        y -= 20

    c.showPage()
    c.save()
    buffer.seek(0)

    return send_file(buffer, mimetype='application/pdf',
                     as_attachment=True, download_name='arthro_insight_results.pdf')

def export_png(results, input_data):
    buffer = io.BytesIO()
    create_results_chart(results)
    plt.savefig(buffer, format='png')
    buffer.seek(0)

    return send_file(buffer, mimetype='image/png',
                     as_attachment=True, download_name='arthro_insight_results.png')

def create_results_chart(results):
    fig, ax = plt.subplots(figsize=(10, 6))

    ax.bar(results.keys(), results.values(), color='skyblue')
    ax.set_xlabel('Models')
    ax.set_ylabel('Likelihood of Joint Replacement (%)')
    ax.set_title('Model Predictions')

    plt.xticks(rotation=45)
    plt.tight_layout()

if __name__ == '__main__':
    app.run(debug=False)  # Set debug to False in production
    app.run(host='0.0.0.0', port=8080, threaded=False)  # Disable threading
